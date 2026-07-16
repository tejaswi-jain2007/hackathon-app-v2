import os
import re
import psycopg2
from werkzeug.security import generate_password_hash
from dotenv import load_dotenv
from datetime import datetime, timezone
import subprocess

# Load dotenv from current backend directory
backend_dir = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(backend_dir, ".env"))

# List of mentors from the jury panel photos
mentors_data = [
    {"name": "Mr. Gaurav Sharma", "designation": "FULL-STACK DEVELOPER", "company": "Tribute technology"},
    {"name": "Ms. Sakshi Prajapat", "designation": "SOFTWARE DEVELOPER", "company": "Thoughttwin"},
    {"name": "Mr. Rakshit Sharma", "designation": "AUTOMATION SPECIALIST", "company": "Walkover"},
    {"name": "Ms. Yogita Patidar", "designation": "SYSTEM ENGINEER", "company": "TCS"},
    {"name": "Mr. Kunal Thakur", "designation": "JUNIOR DOTNET DEVELOPER", "company": "Workpulse Solutions"},
    {"name": "Ms. Aarti Jethhaniya", "designation": "FOUNDER", "company": "Webstarts.in"},
    {"name": "Mr. Tarun Kumar Makode", "designation": "AI ENGINEER", "company": "NeuratantraAI"},
    {"name": "Mr. Ajit Rawat", "designation": "AI/ML DEVELOPER", "company": "MoreYeahs"},
    {"name": "Mr. Yogesh Panchole", "designation": "SOFTWARE ENGINEER", "company": "Neuratantra AI"},
    {"name": "Mr. Vaibhav Jain", "designation": "FOUNDER", "company": "Neuratantra AI"}
]

# Clean name and generate credentials
mentors_credentials = []
for mentor in mentors_data:
    raw_name = mentor["name"]
    # Strip "Mr. " or "Ms. " prefix
    clean_name = raw_name.replace("Mr. ", "").replace("Ms. ", "").strip()
    # Split first/last
    parts = clean_name.split()
    first_name = parts[0] if len(parts) > 0 else "Mentor"
    last_name = parts[1] if len(parts) > 1 else ""
    
    # Generate email
    email_name = f"{first_name.lower()}.{last_name.lower()}" if last_name else first_name.lower()
    email = f"mentor.{email_name}@iist.com"
    
    # Generate password
    password = f"Mentor2026@{first_name}"
    
    mentors_credentials.append({
        "raw_name": raw_name,
        "clean_name": clean_name,
        "email": email,
        "password": password,
        "designation": mentor["designation"],
        "company": mentor["company"]
    })

# Connect to database
db_url = os.getenv("NEON_DB_URL") or os.getenv("DATABASE_URL")
if not db_url:
    print("Error: DATABASE_URL or NEON_DB_URL not set in environment.")
    exit(1)

print("Connecting to Neon Database...")
conn = psycopg2.connect(db_url)
cursor = conn.cursor()

print("Registering mentors...")
registered_count = 0
for mentor in mentors_credentials:
    name = mentor["clean_name"]
    email = mentor["email"]
    password = mentor["password"]
    hashed = generate_password_hash(password)
    created_at = datetime.now(timezone.utc).isoformat()
    
    try:
        # Check if email already exists
        cursor.execute("SELECT id FROM users WHERE email = %s", (email,))
        existing = cursor.fetchone()
        if existing:
            print(f"  Mentor '{name}' ({email}) is already registered. Updating password...")
            cursor.execute(
                "UPDATE users SET name = %s, password_hash = %s WHERE email = %s",
                (name, hashed, email)
            )
        else:
            cursor.execute(
                """
                INSERT INTO users (name, email, password_hash, role, created_at)
                VALUES (%s, %s, %s, 'mentor', %s)
                """,
                (name, email, hashed, created_at)
            )
            print(f"  Successfully registered mentor: {name} ({email})")
            registered_count += 1
    except Exception as e:
        print(f"  Error registering mentor '{name}': {e}")
        conn.rollback()

conn.commit()
conn.close()
print(f"Database registration completed! ({registered_count} new entries)")

# Create Docx report
docx_path = r"c:\Users\tejas\Downloads\Mentor_Credentials_Internal_Hackathon_2026.docx"
print("Attempting to write docx report...")

# Install python-docx if not already installed (usually should be since we installed it earlier)
try:
    import docx
except ImportError:
    print("python-docx not found. Installing via pip...")
    subprocess.check_call(["pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Add styled Title
title = doc.add_paragraph()
title_run = title.add_run("Mentor Panel Credentials")
title_run.font.name = 'Arial'
title_run.font.size = Pt(24)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Dark slate
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
sub_run = subtitle.add_run("Internal Hackathon 2026 | 17-18 July 2026")
sub_run.font.name = 'Arial'
sub_run.font.size = Pt(12)
sub_run.font.italic = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("\nThis document contains the automatically generated credentials for the Hackathon Mentor panel. The Admin can assign teams to these mentors via the Admin Panel under the 'Mentors' tab.\n")

# Add Table
table = doc.add_table(rows=1, cols=5)
table.style = 'Light Shading Accent 1'

hdr_cells = table.rows[0].cells
headers = ["S.No", "Mentor Name", "Affiliation", "Portal Email ID", "Default Password"]
for i, name in enumerate(headers):
    hdr_cells[i].text = name
    # Make header bold
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

for idx, mentor in enumerate(mentors_credentials, 1):
    row_cells = table.add_row().cells
    row_cells[0].text = str(idx)
    row_cells[1].text = mentor["clean_name"]
    row_cells[2].text = f"{mentor['designation']}, {mentor['company']}"
    row_cells[3].text = mentor["email"]
    row_cells[4].text = mentor["password"]

doc.save(docx_path)
print(f"Docx report created successfully at: {docx_path}")

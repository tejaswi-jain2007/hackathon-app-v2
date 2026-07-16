import os
import re
import psycopg2
from werkzeug.security import generate_password_hash
from dotenv import load_dotenv
from datetime import datetime, timezone
import subprocess

# Load dotenv from current backend directory
backend_dir = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(backend_dir, ".env"))

# List of judges from the jury panel photos
judges_data = [
    {"name": "Mr. Kavish Mehta", "designation": "Sr. Member Tech.", "company": "Oracle"},
    {"name": "Mr. Pradeep Raghuwanshi", "designation": "Senior Consultant", "company": "TCS"},
    {"name": "Mr. Pawan Patidar", "designation": "Lead Software Engineer", "company": "Impetus"},
    {"name": "Mr. Darshan Joshi", "designation": "Associate Consultant", "company": "TCS"},
    {"name": "Mr. Ashish Dave", "designation": "Lead Auditor", "company": "Infobeans"},
    {"name": "Mr. Pavan Jain", "designation": "Lead Software Engineer", "company": "Lirisoft"},
    {"name": "Mr. Jitendra Gupta", "designation": "Sr. Software Engineer", "company": "Microsoft"},
    {"name": "Mr. Suyash Heda", "designation": "Sr. Software Engineer", "company": "Microsoft"},
    {"name": "Mr. Nitish Tongia", "designation": "Software Engineer II", "company": "Oracle"},
    {"name": "Mr. Rasesh Tongia", "designation": "Sr. Member Tech.", "company": "Oracle"},
    {"name": "Mr. Praveen Kushwah", "designation": "Technical Lead", "company": "TCS"},
    {"name": "Mr. Nandkishore Patidar", "designation": "Technical Architect", "company": "TCS"}
]

# Clean name and generate credentials
judges_credentials = []
for judge in judges_data:
    raw_name = judge["name"]
    # Strip "Mr. " prefix
    clean_name = raw_name.replace("Mr. ", "").strip()
    # Split first/last
    parts = clean_name.split()
    first_name = parts[0] if len(parts) > 0 else "Judge"
    last_name = parts[1] if len(parts) > 1 else ""
    
    # Generate email
    email_name = f"{first_name.lower()}.{last_name.lower()}" if last_name else first_name.lower()
    email = f"judge.{email_name}@iist.com"
    
    # Generate password
    password = f"Jury2026@{first_name}"
    
    judges_credentials.append({
        "raw_name": raw_name,
        "clean_name": clean_name,
        "email": email,
        "password": password,
        "designation": judge["designation"],
        "company": judge["company"]
    })

# Connect to database
db_url = os.getenv("NEON_DB_URL") or os.getenv("DATABASE_URL")
if not db_url:
    print("Error: DATABASE_URL or NEON_DB_URL not set in environment.")
    exit(1)

print("Connecting to Neon Database...")
conn = psycopg2.connect(db_url)
cursor = conn.cursor()

print("Registering judges...")
registered_count = 0
for judge in judges_credentials:
    name = judge["clean_name"]
    email = judge["email"]
    password = judge["password"]
    hashed = generate_password_hash(password)
    created_at = datetime.now(timezone.utc).isoformat()
    
    try:
        # Check if email already exists
        cursor.execute("SELECT id FROM users WHERE email = %s", (email,))
        existing = cursor.fetchone()
        if existing:
            print(f"  Judge '{name}' ({email}) is already registered. Updating password...")
            cursor.execute(
                "UPDATE users SET name = %s, password_hash = %s WHERE email = %s",
                (name, hashed, email)
            )
        else:
            cursor.execute(
                """
                INSERT INTO users (name, email, password_hash, role, created_at)
                VALUES (%s, %s, %s, 'judge', %s)
                """,
                (name, email, hashed, created_at)
            )
            print(f"  Successfully registered judge: {name} ({email})")
            registered_count += 1
    except Exception as e:
        print(f"  Error registering judge '{name}': {e}")
        conn.rollback()

conn.commit()
conn.close()
print(f"Database registration completed! ({registered_count} new entries)")

# Create Docx report
docx_path = r"c:\Users\tejas\Downloads\Jury_Credentials_Internal_Hackathon_2026.docx"
print("Attempting to write docx report...")

# Install python-docx if not already installed
try:
    import docx
except ImportError:
    print("python-docx not found. Installing via pip...")
    subprocess.check_call(["pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Add styled Title
title = doc.add_paragraph()
title_run = title.add_run("Jury Panel Credentials")
title_run.font.name = 'Arial'
title_run.font.size = Pt(24)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Dark slate
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
sub_run = subtitle.add_run("Internal Hackathon 2026 | 17-18 July 2026")
sub_run.font.name = 'Arial'
sub_run.font.size = Pt(12)
sub_run.font.italic = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("\nThis document contains the automatically generated credentials for the Hackathon Jury panel. The Admin can assign teams to these judges via the Admin Panel under the 'Judges' tab.\n")

# Add Table
table = doc.add_table(rows=1, cols=5)
table.style = 'Light Shading Accent 1'

hdr_cells = table.rows[0].cells
headers = ["S.No", "Judge Name", "Affiliation", "Portal Email ID", "Default Password"]
for i, name in enumerate(headers):
    hdr_cells[i].text = name
    # Make header bold
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

for idx, judge in enumerate(judges_credentials, 1):
    row_cells = table.add_row().cells
    row_cells[0].text = str(idx)
    row_cells[1].text = judge["clean_name"]
    row_cells[2].text = f"{judge['designation']}, {judge['company']}"
    row_cells[3].text = judge["email"]
    row_cells[4].text = judge["password"]

doc.save(docx_path)
print(f"Docx report created successfully at: {docx_path}")
